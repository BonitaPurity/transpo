from docx import Document

agreements = {
    'docs/TRANSPO_HUB_NDA.docx': {
        'title': 'NON-DISCLOSURE AGREEMENT (TRANSPO HUB)',
        'sections': [
            ('Parties', 'This Non-Disclosure Agreement ("Agreement") is entered into by and between Person 1 ("Disclosing Party") and Person 2 ("Receiving Party").'),
            ('1. Purpose', 'The purpose of this Agreement is to permit Person 2 to view and use the TRANSPO HUB software and related confidential information exclusively for the university final-year project demonstration, while protecting the confidentiality of that information.'),
            ('2. Confidential Information', 'Confidential Information includes, without limitation, source code, database structure, architecture diagrams, implementation details, credentials, roadmaps, test cases, deployment procedures, and any other technical or business information related to TRANSPO HUB that is not publicly available.'),
            ('3. Obligations of Receiving Party', 'The Receiving Party agrees to: (a) keep all Confidential Information strictly confidential; (b) use the Confidential Information solely for the final-year project demonstration; (c) not disclose or distribute the Confidential Information to any third party; (d) protect the Confidential Information with commercially reasonable care; and (e) return or destroy all copies of the Confidential Information upon written request or upon completion of the demonstration.'),
            ('4. Term', 'This Agreement remains in effect for five (5) years from the date of disclosure of the Confidential Information, or longer if required by applicable law.'),
            ('5. Exceptions', 'Confidential Information does not include information that: (a) is or becomes publicly known through no breach of this Agreement by the Receiving Party; (b) is independently developed by the Receiving Party without use of the Confidential Information; or (c) is required to be disclosed by law, provided the Receiving Party gives prompt notice to the Disclosing Party to allow for protective measures.'),
            ('6. Remedies', 'The Receiving Party acknowledges that breach of this Agreement may cause irreparable harm to the Disclosing Party and agrees that the Disclosing Party is entitled to seek injunctive relief and/or damages.'),
        ],
    },
    'docs/TRANSPO_HUB_COPYRIGHT_IP.docx': {
        'title': 'COPYRIGHT AND INTELLECTUAL PROPERTY ASSIGNMENT (TRANSPO HUB)',
        'sections': [
            ('Grant of Ownership', 'Person 2 acknowledges that all copyright and intellectual property rights in the TRANSPO HUB software (including source code, documentation, designs, and any enhancements) are owned exclusively by Person 1.'),
            ('Assignment', 'Person 2 hereby assigns, transfers, and conveys to Person 1 all rights, title, and interest in any modifications, additions, or contributions made by Person 2 to TRANSPO HUB, including all copyright and moral rights in such contributions.'),
            ('Moral Rights Waiver', 'To the fullest extent permitted by law, Person 2 waives any moral rights, including rights of attribution and integrity, with respect to the TRANSPO HUB software and any derivatives.'),
            ('No Retained Rights', 'Person 2 acknowledges that no ownership interest is retained; Person 2 is granted only the limited usage permission set forth in the separate Educational License below.'),
        ],
    },
    'docs/TRANSPO_HUB_EDUCATIONAL_LICENSE.docx': {
        'title': 'EDUCATIONAL LICENSE (TRANSPO HUB)',
        'sections': [
            ('Grant', 'Person 1 grants Person 2 a non-transferable, non-exclusive, revocable license to use the TRANSPO HUB software solely for the purpose of preparing and presenting the university final-year project demonstration.'),
            ('Permitted Use', 'Under this license, Person 2 may: (a) run and demonstrate the software in a classroom or evaluation setting; (b) make minor modifications necessary for the demonstration; and (c) reference the software in project documentation that is limited to the final-year project evaluation context.'),
            ('Restrictions', 'Person 2 may not: (a) distribute, publish, sublicense, or otherwise make the software available to third parties; (b) use the software for commercial, production, or any non-educational purposes; (c) deploy the software as a service or product outside of the demonstration context.'),
            ('Term', 'This license automatically terminates upon completion of the final-year project demonstration, or earlier if revoked in writing by Person 1. Upon termination, Person 2 must delete all copies of the software and related materials.'),
            ('Acknowledgement', 'Person 2 acknowledges that all rights in TRANSPO HUB remain with Person 1 and that this license does not transfer any ownership rights.'),
        ],
    },
}

for path, data in agreements.items():
    doc = Document()
    doc.add_heading(data['title'], level=1)
    for heading, body in data['sections']:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    # Ensure directory exists
    import os
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)

print('Created documents:', ', '.join(agreements.keys()))
